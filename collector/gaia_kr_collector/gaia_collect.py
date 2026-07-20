#!/usr/bin/env python3
"""
Gaia Korean Dataset Collector
- Excludes SEC / DART / AI Hub OCR by design.
- Collects Korean-heavy, text-bearing PDF/XLSX/CSV/XML/JSON/TXT/DOCX/DOC files.
- HWP/HWPX files are downloaded and converted to DOCX via LibreOffice (libreoffice-h2orestart).
  Conversion failures (~5%) are logged and skipped.
- Quota is tracked per document TYPE (not per source) — run any source as many
  times as needed; duplicates are skipped via manifest.csv.

Usage:
  python gaia_collect.py --config config.yaml --plan
  python gaia_collect.py --config config.yaml --run all
  python gaia_collect.py --config config.yaml --run gov_policy_reports data_go_kr
  python gaia_collect.py --config config.yaml --run all --bg
  python gaia_collect.py --config config.yaml --run all --bg --log /tmp/collect.log

Notes:
  * Some portals require login/API approval or dynamically generated download URLs.
    For those, add concrete seed URLs to config.yaml.
  * Common Crawl extraction downloads WET files and writes Korean text chunks.
  * Check --plan output to see which document type still needs more data, then
    re-run the most productive sources for that type.
"""
from __future__ import annotations

import argparse
import bz2
import csv
import gzip
import hashlib
import json
import os
import random
import re
import shutil
import subprocess
import sys
import threading
import time
import xml.etree.ElementTree as ET
from collections import deque
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional, Set
from urllib.parse import urljoin, urlparse, unquote

import requests
import yaml
from bs4 import BeautifulSoup
from tqdm import tqdm

try:
    from warcio.archiveiterator import ArchiveIterator
except Exception:
    ArchiveIterator = None

SIZE_RE = re.compile(r"^\s*(\d+(?:\.\d+)?)\s*([KMGT]?B)\s*$", re.I)
FILE_EXT_RE = re.compile(r"\.(pdf|csv|xlsx|xls|xml|json|txt|docx|doc)(?:$|[?#])", re.I)
HWP_EXT_RE  = re.compile(r"\.(hwp|hwpx)(?:$|[?#])", re.I)
HANGUL_RE = re.compile(r"[가-힣]")
URL_LIKE_RE = re.compile(r"https?://[^\s\"'<>]+")

# 다운로드 URL의 실제 확장자가 경로가 아니라 쿼리스트링에 있는 경우 (예: lbFileDownload.do?...&flExt=pdf)
KNOWN_DOC_EXTS = {"pdf", "csv", "xlsx", "xls", "xml", "json", "txt", "docx", "doc",
                  "hwp", "hwpx", "odf", "rtf", "ppt", "pptx", "html", "zip"}
QUERY_EXT_RE = re.compile(r"(?:flExt|fileExt|fileExtension|ext|type)=([A-Za-z0-9]{2,5})", re.I)

# Content-Type → 저장할 확장자. None이면 저장하지 않음.
CONTENT_TYPE_EXT: Dict[str, Optional[str]] = {
    "application/pdf":                                                          ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "application/msword":                                                       ".doc",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":       ".xlsx",
    "application/vnd.ms-excel":                                                 ".xls",
    "text/csv":                                                                 ".csv",
    "text/plain":                                                               ".txt",
    "application/json":                                                         None,
    "application/xml":                                                          ".xml",
    "text/xml":                                                                 ".xml",
    "text/html":                                                                ".html",
    "application/rtf":                                                          ".rtf",
    "text/rtf":                                                                 ".rtf",
    "application/vnd.oasis.opendocument.text":                                  ".odf",
    "application/vnd.oasis.opendocument.spreadsheet":                           ".odf",
    "application/vnd.ms-powerpoint":                                            ".ppt",
    "application/vnd.openxmlformats-officedocument.presentationml.presentation": ".pptx",
    "application/vnd.hancom.hwp":                                               ".hwp",
    "application/vnd.hancom.hwpx":                                              ".hwpx",
    # 아래는 저장 불필요 → None
    "image/jpeg": None, "image/png": None, "image/gif": None, "image/webp": None,
    "application/zip": None, "application/x-zip-compressed": None,
    "application/x-rar-compressed": None, "application/octet-stream": None,
    "application/javascript": None,
}

# Keys in config.yaml that are NOT source names.
NON_SOURCE_KEYS = frozenset({
    "root_dir", "user_agent", "request_timeout_sec", "sleep_sec",
    "max_retries", "quotas", "allowed_extensions", "generic_crawl",
    "parallel_workers",
})

# Document type groups and the file extensions that belong to each.
TYPE_GROUPS: Dict[str, Set[str]] = {
    "pdf":          {"pdf"},
    "docx_doc":     {"docx", "doc", "odf", "rtf"},
    "xlsx_xls_csv": {"xlsx", "xls", "csv"},
    "ppt_pptx":     {"ppt", "pptx"},
    "txt":          {"txt", "html"},
}
EXT_TO_GROUP: Dict[str, str] = {
    ext: grp for grp, exts in TYPE_GROUPS.items() for ext in exts
}


def parse_size(s: str) -> int:
    m = SIZE_RE.match(str(s))
    if not m:
        raise ValueError(f"Invalid size: {s}")
    n = float(m.group(1)); unit = m.group(2).upper()
    mult = {"B": 1, "KB": 10**3, "MB": 10**6, "GB": 10**9, "TB": 10**12}[unit]
    return int(n * mult)


def human(n: int) -> str:
    for unit in ["B", "KB", "MB", "GB", "TB"]:
        if n < 1000 or unit == "TB":
            return f"{n:.1f}{unit}" if unit != "B" else f"{n}B"
        n /= 1000
    return str(n)


def mkdir(p: Path) -> None:
    p.mkdir(parents=True, exist_ok=True)


def sha1_text(s: str) -> str:
    return hashlib.sha1(s.encode("utf-8", errors="ignore")).hexdigest()


def safe_name_from_url(url: str, default_ext: str = ".bin") -> str:
    parsed = urlparse(url)
    path_name = Path(unquote(parsed.path)).name
    path_ext = Path(path_name).suffix.lower().lstrip(".")
    if path_name and path_ext in KNOWN_DOC_EXTS:
        name = path_name
    else:
        # 경로에 알려진 문서 확장자가 없으면 쿼리스트링에서 실제 확장자를 찾는다
        # (예: .../download.do?fileId=1&flExt=pdf 같은 URL 패턴)
        m = QUERY_EXT_RE.search(parsed.query)
        q_ext = m.group(1).lower() if m else ""
        ext = f".{q_ext}" if q_ext in KNOWN_DOC_EXTS else default_ext
        name = sha1_text(url)[:16] + ext
    name = re.sub(r"[^0-9A-Za-z가-힣._-]+", "_", name)
    return name[:180]


@dataclass
class Cfg:
    raw: dict
    root: Path
    ua: str
    sleep: float
    timeout: int
    retries: int
    allowed_exts: Set[str]

    @classmethod
    def load(cls, path: str) -> "Cfg":
        raw = yaml.safe_load(Path(path).read_text(encoding="utf-8"))
        return cls(
            raw=raw,
            root=Path(raw.get("root_dir", "../../gaia_web_dataset")),
            ua=raw.get("user_agent", "gaia-test-collector/1.0"),
            sleep=float(raw.get("sleep_sec", 1.0)),
            timeout=int(raw.get("request_timeout_sec", 30)),
            retries=int(raw.get("max_retries", 3)),
            allowed_exts={e.lower() for e in raw.get("allowed_extensions", [])},
        )

    @property
    def total_quota(self) -> int:
        return parse_size(self.raw.get("quotas", {}).get("total", "200GB"))

    @property
    def type_quotas(self) -> Dict[str, int]:
        return {k: parse_size(v) for k, v in self.raw.get("quotas", {}).get("by_type", {}).items()}

    @property
    def lang_caps(self) -> Dict[str, float]:
        """타입별 영어 소스 허용 비율 (예: docx_doc: 0.3 → 영어는 해당 타입 quota의 30%까지만)."""
        return {k: float(v) for k, v in self.raw.get("quotas", {}).get("lang_cap", {}).items()}

    @property
    def parallel_workers(self) -> int:
        return int(self.raw.get("parallel_workers", 1))

    def all_sources(self) -> List[str]:
        return [k for k in self.raw if k not in NON_SOURCE_KEYS]

    def source_lang(self, source: str) -> str:
        return str(self.raw.get(source, {}).get("lang", "ko")).lower()


class Collector:
    def __init__(self, cfg: Cfg):
        self.cfg = cfg
        self._local = threading.local()   # per-thread HTTP session
        self._lock = threading.Lock()     # _type_bytes + manifest.csv 보호
        self._lo_lock = threading.Lock()  # LibreOffice 직렬화
        mkdir(cfg.root)
        self.manifest_path = cfg.root / "manifest.csv"
        if not self.manifest_path.exists():
            with self.manifest_path.open("w", encoding="utf-8", newline="") as f:
                csv.writer(f).writerow(["ts", "source", "url", "path", "bytes", "status"])
        if cfg.allowed_exts:
            exts = "|".join(re.escape(e.lstrip(".")) for e in sorted(cfg.allowed_exts))
            self._file_ext_re = re.compile(rf"\.({exts})(?:$|[?#])", re.I)
        else:
            self._file_ext_re = FILE_EXT_RE
        # LibreOffice 가용 여부 (HWP→DOCX 변환에 사용)
        self._libreoffice_ok: bool = shutil.which("libreoffice") is not None
        if self._libreoffice_ok:
            print("[INFO] LibreOffice 감지됨 — HWP/HWPX 파일을 DOCX로 변환합니다 (원본 보존).")
        # Initialize per-type byte counters from manifest history.
        self._type_bytes: Dict[str, int] = self._init_type_bytes()
        self._type_en_bytes: Dict[str, int] = self._init_type_bytes(lang="en")
        # 이미 방문한 URL 세트 (페이지 + 파일) — 재시작 시 재크롤 방지
        self._visited_path = cfg.root / "visited_pages.txt"
        self._visited: Set[str] = self._load_visited()

    def _load_visited(self) -> Set[str]:
        """visited_pages.txt + manifest.csv URL을 합쳐 seen 초기 세트 반환."""
        visited: Set[str] = set()
        if self._visited_path.exists():
            with self._visited_path.open(encoding="utf-8", errors="ignore") as f:
                for line in f:
                    u = line.strip()
                    if u:
                        visited.add(u)
        if self.manifest_path.exists():
            with self.manifest_path.open(encoding="utf-8", newline="", errors="ignore") as f:
                for row in csv.DictReader(f):
                    u = row.get("url", "").strip()
                    if u:
                        visited.add(u)
        return visited

    def _mark_visited(self, url: str) -> None:
        """페이지 URL을 visited_pages.txt에 기록."""
        with self._lock:
            self._visited.add(url)
            with self._visited_path.open("a", encoding="utf-8") as f:
                f.write(url + "\n")

    # ------------------------------------------------------------------
    # Type-quota helpers
    # ------------------------------------------------------------------

    def _init_type_bytes(self, lang: Optional[str] = None) -> Dict[str, int]:
        """매니페스트 이력에서 타입별 누적 바이트 집계. lang="en"이면 영어 소스분만 집계."""
        counts: Dict[str, int] = {g: 0 for g in TYPE_GROUPS}
        if not self.manifest_path.exists():
            return counts
        with self.manifest_path.open(encoding="utf-8", newline="") as f:
            for row in csv.DictReader(f):
                if row.get("status") in ("downloaded", "docx_batch", "ko_chunk"):
                    if lang is not None and self.cfg.source_lang(row.get("source", "")) != lang:
                        continue
                    ext = Path(row.get("path", "")).suffix.lower().lstrip(".")
                    grp = EXT_TO_GROUP.get(ext)
                    if grp:
                        counts[grp] = counts.get(grp, 0) + int(row.get("bytes", 0) or 0)
        return counts

    def _session(self) -> requests.Session:
        if not hasattr(self._local, "session"):
            s = requests.Session()
            s.headers.update({"User-Agent": self.cfg.ua})
            self._local.session = s
        return self._local.session

    def _ext_group(self, url_or_path: str) -> Optional[str]:
        ext = Path(safe_name_from_url(url_or_path)).suffix.lower().lstrip(".")
        return EXT_TO_GROUP.get(ext)

    def total_collected(self) -> int:
        with self._lock:
            return sum(self._type_bytes.values())

    def under_type_quota(self, url_or_path: str, source: Optional[str] = None) -> bool:
        grp = self._ext_group(url_or_path)
        if grp is None:
            return True
        limit = self.cfg.type_quotas.get(grp, 0)
        if limit == 0:
            return True
        with self._lock:
            if self._type_bytes.get(grp, 0) >= limit:
                return False
            if source is not None and self.cfg.source_lang(source) == "en":
                cap = self.cfg.lang_caps.get(grp)
                if cap is not None and self._type_en_bytes.get(grp, 0) >= limit * cap:
                    return False  # 영어 소스는 타입별 허용 비율(예: 30%)까지만
            return True

    def any_quota_remaining(self) -> bool:
        """타입별 쿼터가 하나라도 남아있으면 True.
        전체 합계(quotas.total)는 개별 타입 쿼터의 합과 같아야 정상이므로 별도로
        게이트하지 않는다 — 한 타입이 버그로 초과 수집되어도 다른 타입의 목표
        달성을 막지 않도록 하기 위함 (2026-07-03: txt 초과로 인한 조기 종료 버그)."""
        return bool(self.lagging_types())

    def lagging_types(self) -> set:
        """쿼터 미달 타입 집합 반환."""
        with self._lock:
            return {grp for grp, limit in self.cfg.type_quotas.items()
                    if self._type_bytes.get(grp, 0) < limit}

    def _type_remain(self, url_or_path: str, source: Optional[str] = None) -> int:
        grp = self._ext_group(url_or_path)
        if grp is None and HWP_EXT_RE.search(url_or_path):
            grp = "docx_doc"
        if grp is None:
            return self.cfg.total_quota  # 타입을 알 수 없는 경우 사실상 무제한
        limit = self.cfg.type_quotas.get(grp, 0)
        with self._lock:
            type_cur = self._type_bytes.get(grp, 0)
            en_cur = self._type_en_bytes.get(grp, 0)
        remain = max(0, limit - type_cur) if limit else self.cfg.total_quota
        if source is not None and self.cfg.source_lang(source) == "en":
            cap = self.cfg.lang_caps.get(grp)
            if cap is not None and limit:
                en_remain = max(0, int(limit * cap) - en_cur)
                remain = min(remain, en_remain)
        return remain

    def convert_hwp(self, hwp_path: Path) -> Optional[Path]:
        """HWP/HWPX → DOCX 변환. 원본은 보존하고 변환된 DOCX 경로를 반환."""
        out_dir = hwp_path.parent
        docx_path = out_dir / (hwp_path.stem + ".docx")
        if docx_path.exists() and docx_path.stat().st_size > 0:
            return docx_path
        with self._lo_lock:  # LibreOffice는 동시 실행 불가 — 직렬화
            if docx_path.exists() and docx_path.stat().st_size > 0:
                return docx_path
            try:
                result = subprocess.run(
                    ["libreoffice", "--headless", "--convert-to", "docx",
                     "--outdir", str(out_dir), str(hwp_path)],
                    capture_output=True, timeout=120,
                )
                if result.returncode == 0 and docx_path.exists() and docx_path.stat().st_size > 0:
                    return docx_path
            except (subprocess.TimeoutExpired, FileNotFoundError):
                pass
        return None

    # ------------------------------------------------------------------
    # I/O helpers
    # ------------------------------------------------------------------

    def log(self, source: str, url: str, path: Path, nbytes: int, status: str) -> None:
        with self._lock:
            with self.manifest_path.open("a", encoding="utf-8", newline="") as f:
                csv.writer(f).writerow([int(time.time()), source, url, str(path), nbytes, status])
            if status in ("downloaded", "docx_batch", "ko_chunk") and nbytes > 0:
                grp = EXT_TO_GROUP.get(path.suffix.lower().lstrip("."))
                if grp:
                    self._type_bytes[grp] = self._type_bytes.get(grp, 0) + nbytes
                    if self.cfg.source_lang(source) == "en":
                        self._type_en_bytes[grp] = self._type_en_bytes.get(grp, 0) + nbytes

    def get(self, url: str, stream: bool = False) -> Optional[requests.Response]:
        for attempt in range(self.cfg.retries):
            try:
                r = self._session().get(url, timeout=self.cfg.timeout, stream=stream, allow_redirects=True)
                if r.status_code in (429, 500, 502, 503, 504):
                    time.sleep(self.cfg.sleep * (2 ** attempt + random.random()))
                    continue
                if r.status_code >= 400:
                    return None
                return r
            except requests.RequestException:
                time.sleep(self.cfg.sleep * (2 ** attempt + random.random()))
        return None

    @staticmethod
    def _cd_filename(headers) -> Optional[str]:
        """Content-Disposition 헤더에서 실제 파일명 추출."""
        cd = headers.get("content-disposition", "")
        if not cd:
            return None
        # RFC 5987: filename*=UTF-8''%EC%9D%B4...
        m = re.search(r"filename\*\s*=\s*[A-Za-z0-9-]*''([^;\s]+)", cd, re.I)
        if m:
            try:
                return unquote(m.group(1), encoding="utf-8")
            except Exception:
                pass
        # 일반 filename="..."
        m = re.search(r'filename\s*=\s*"([^"]+)"', cd, re.I)
        if m:
            raw = m.group(1)
            try:
                raw = raw.encode("latin-1").decode("utf-8")  # 한국 사이트 EUC-KR→UTF-8 복구
            except Exception:
                pass
            # RFC 5987이 아니라 filename="%EC%9D%B4..." 처럼 따옴표 안에 퍼센트
            # 인코딩된 바이트를 그대로 넣어 보내는 사이트가 많다 (2026-07-20:
            # kdischool_eng 등 파일명 깨짐 버그). 위 latin-1 복구는 순수 ASCII인
            # 퍼센트 인코딩 문자열엔 아무 효과가 없으므로 별도로 unquote한다.
            if re.search(r"%[0-9A-Fa-f]{2}", raw):
                try:
                    raw = unquote(raw, encoding="utf-8", errors="strict")
                except Exception:
                    pass
            return raw
        m = re.search(r"filename\s*=\s*([^;\s\"']+)", cd, re.I)
        if m:
            return unquote(m.group(1))
        return None

    def download(self, source: str, url: str, out_dir: Path) -> bool:
        mkdir(out_dir)
        if not self.any_quota_remaining():
            return False
        name = safe_name_from_url(url)
        is_hwp = bool(HWP_EXT_RE.search(name))
        if is_hwp and not self._libreoffice_ok:
            return False
        if not self.under_type_quota("dummy.docx" if is_hwp else url, source=source):
            return False
        out = out_dir / name
        if out.exists() and out.stat().st_size > 0:
            return True
        tmp = out.with_suffix(out.suffix + ".part")
        r = self.get(url, stream=True)
        if not r:
            self.log(source, url, out, 0, "http_error")
            return False

        # 1) Content-Disposition에서 실제 파일명 우선 추출
        #    한국 공공사이트는 .do URL이라도 여기에 실제 확장자가 있음
        cd_name = self._cd_filename(r.headers)
        if cd_name:
            cd_ext = Path(cd_name).suffix.lower()
            if cd_ext in self.cfg.allowed_exts or HWP_EXT_RE.search(cd_name):
                safe_cd = re.sub(r"[^0-9A-Za-z가-힣._-]+", "_", cd_name)[:180]
                name = safe_cd
                out = out_dir / name
                tmp = out.with_suffix(out.suffix + ".part")
                if out.exists() and out.stat().st_size > 0:
                    return True
                is_hwp = bool(HWP_EXT_RE.search(name))

        # 2) Content-Type으로 보조 확인
        ctype = r.headers.get("content-type", "").split(";")[0].strip().lower()
        if ctype and not is_hwp:
            ct_ext = CONTENT_TYPE_EXT.get(ctype)
            if ct_ext is None and ctype in CONTENT_TYPE_EXT:
                # octet-stream은 CD에서 이미 처리 시도 — 여기까지 오면 허용 확장자 아님
                if ctype != "application/octet-stream":
                    self.log(source, url, out, 0, "skip_content_type")
                    return False
            if ct_ext and name.endswith(".bin"):
                name = name[:-4] + ct_ext
                out = out_dir / name
                tmp = out.with_suffix(out.suffix + ".part")
                if out.exists() and out.stat().st_size > 0:
                    return True
                is_hwp = bool(HWP_EXT_RE.search(name))

        # 3) 확장자를 끝내 알 수 없으면 스킵
        if name.endswith(".bin"):
            self.log(source, url, out, 0, "skip_unknown_type")
            return False

        total = int(r.headers.get("content-length", 0) or 0)
        remain = self._type_remain("dummy.docx" if is_hwp else name, source=source)
        if total and total > remain:
            self.log(source, url, out, 0, "skip_over_quota")
            return False
        n = 0
        with tmp.open("wb") as f:
            for chunk in r.iter_content(chunk_size=1024 * 1024):
                if not chunk:
                    continue
                f.write(chunk); n += len(chunk)
                if n > remain:
                    break
        if n == 0:
            tmp.unlink(missing_ok=True)
            self.log(source, url, out, 0, "empty")
            return False
        tmp.rename(out)
        if is_hwp:
            docx = self.convert_hwp(out)
            if docx:
                self.log(source, url, docx, docx.stat().st_size, "downloaded")
            else:
                self.log(source, url, out, 0, "convert_failed")
        else:
            self.log(source, url, out, n, "downloaded")
        time.sleep(self.cfg.sleep)
        return True

    def is_allowed_file_url(self, url: str) -> bool:
        lower = url.lower()
        if self._file_ext_re.search(lower):
            return True
        if self._libreoffice_ok and HWP_EXT_RE.search(lower):
            return True  # HWP/HWPX → LibreOffice로 DOCX 변환
        return any(x in lower for x in ["download", "filedown", "attach", "atchfile", "getfile", "downfile"])

    @staticmethod
    def _clean_url(url: str) -> str:
        return url.replace("[", "%5B").replace("]", "%5D")

    def extract_links(self, base_url: str, html: str) -> List[str]:
        soup = BeautifulSoup(html, "lxml")
        links = []
        for tag in soup.find_all(["a", "link", "script"]):
            href = tag.get("href") or tag.get("src")
            if href:
                try:
                    links.append(urljoin(base_url, self._clean_url(href)))
                except ValueError:
                    pass
        for m in URL_LIKE_RE.finditer(html):
            links.append(self._clean_url(m.group(0)))
        out = []
        seen: Set[str] = set()
        for u in links:
            try:
                p = urlparse(u)
            except ValueError:
                continue
            if p.scheme not in ("http", "https"):
                continue
            u = u.split("#")[0]
            if u not in seen:
                seen.add(u); out.append(u)
        return out

    # ------------------------------------------------------------------
    # Source runners
    # ------------------------------------------------------------------

    def crawl_generic(self, source: str, seed_urls: List[str]) -> None:
        source_dir = self.cfg.root / source
        mkdir(source_dir)
        settings = self.cfg.raw.get("generic_crawl", {})
        max_pages = int(settings.get("max_pages_per_site", 5000))
        max_depth = int(settings.get("max_depth", 4))
        # 이전 실행에서 방문한 URL로 seen 초기화 → 재시작 시 재크롤 방지.
        # 단, seed_urls(게시판 목록 진입점)는 매 실행 제외 대상에서 뺀다 — 안 그러면
        # 예전에 한 번이라도 방문한 소스는 seed 자체가 seen에 남아 프론티어가 항상
        # 비고, 새 글이 올라와도 영구히 0page로 멈춘다 (2026-07-16: nhis_stats 등
        # 재크롤 영구 정지 버그). 목록 페이지를 다시 훑어도 이미 받은 개별 파일/
        # 하위 페이지는 seen에 남아있으므로 재다운로드되지 않는다.
        seen: Set[str] = set(self._visited) - set(seed_urls)
        # 미방문 프론티어(큐)를 파일로 영속화 — seed_urls는 첫 방문 즉시 visited
        # 처리되므로, 프론티어를 저장해두지 않으면 재시작 시 큐가 비어 그 소스는
        # 영구히 0페이지로 멈춘다 (2026-07-08: kostat_eng 등 재시작 후 크롤 유실 버그).
        frontier_path = source_dir / ".frontier.json"
        q: deque = deque()
        if frontier_path.exists():
            try:
                saved = json.loads(frontier_path.read_text())
                q = deque((u, d) for u, d in saved if u not in seen)
            except (json.JSONDecodeError, OSError, ValueError):
                q = deque()
        if not q:
            q = deque((u, 0) for u in seed_urls if u not in seen)
        pages = 0
        seed_domains = {urlparse(self._clean_url(u)).netloc for u in seed_urls}

        def _save_frontier() -> None:
            try:
                tmp = frontier_path.with_suffix(".json.tmp")
                tmp.write_text(json.dumps(list(q)))
                tmp.replace(frontier_path)
            except OSError:
                pass

        # 프론티어를 페이지마다 통째로 다시 직렬화하면, 사이트에 따라 큐가
        # 수십만~백만 건까지 불어났을 때 페이지당 비용이 계속 커져 사실상
        # 크롤이 멈춘 것처럼 느려진다 (2026-07-20: riss_theses 큐 100만건에서
        # 페이지당 20초+까지 느려짐). 시간 간격을 두고 저장한다 — 중단 시에는
        # finally에서 마지막으로 한 번 더 저장하므로 재시작 유실은 없다.
        FRONTIER_SAVE_INTERVAL_SEC = 30
        last_frontier_save = time.monotonic()

        pbar = tqdm(desc=source, unit="page")
        try:
            while q and self.any_quota_remaining() and pages < max_pages:
                url, depth = q.popleft()
                if url in seen:
                    continue
                seen.add(url)
                if self.is_allowed_file_url(url):
                    self.download(source, url, source_dir / "files")
                    continue
                if depth > max_depth:
                    continue
                try:
                    if urlparse(url).netloc not in seed_domains:
                        continue
                except ValueError:
                    continue
                r = self.get(url)
                if not r:
                    continue
                ctype = r.headers.get("content-type", "")
                text = r.text if "text" in ctype or "html" in ctype or not ctype else ""
                if not text:
                    continue
                pages += 1; pbar.update(1)
                self._mark_visited(url)   # 페이지 방문 기록
                links = self.extract_links(url, text)
                for link in links:
                    if link in seen:
                        continue
                    if self.is_allowed_file_url(link):
                        q.appendleft((link, depth + 1))
                    else:
                        try:
                            if depth + 1 <= max_depth and urlparse(link).netloc in seed_domains:
                                q.append((link, depth + 1))
                        except ValueError:
                            pass
                now = time.monotonic()
                if now - last_frontier_save >= FRONTIER_SAVE_INTERVAL_SEC:
                    _save_frontier()
                    last_frontier_save = now
                time.sleep(self.cfg.sleep)
        finally:
            _save_frontier()
            pbar.close()

    def run_kowiki(self) -> None:
        source = "kowiki_knowledge"
        raw_dir = self.cfg.root / source / "raw"
        docx_dir = self.cfg.root / source / "docx"
        mkdir(raw_dir); mkdir(docx_dir)

        for url in self.cfg.raw.get(source, {}).get("urls", []):
            if not self.any_quota_remaining():
                break
            self.download(source, url, raw_dir)
            if "pages-articles-multistream.xml.bz2" in url:
                local = raw_dir / safe_name_from_url(url)
                if local.exists():
                    self._kowiki_to_docx(source, local, docx_dir)

    @staticmethod
    def _strip_wiki_markup(text: str) -> str:
        for _ in range(5):
            text, n = re.subn(r'\{\{[^{}]*\}\}', '', text)
            if n == 0:
                break
        text = re.sub(r'\[\[(?:파일|File|Image|그림|Category|분류)[^\]]*\]\]', '', text, flags=re.I)
        text = re.sub(r'\[\[(?:[^|\]]*\|)?([^\]]*)\]\]', r'\1', text)
        text = re.sub(r'\[https?://\S+\s+([^\]]+)\]', r'\1', text)
        text = re.sub(r'\[https?://\S+\]', '', text)
        text = re.sub(r"'{2,3}", '', text)
        text = re.sub(r'={2,6}\s*(.+?)\s*={2,6}', r'\n\1\n', text)
        text = re.sub(r'<ref[^>]*>.*?</ref>', '', text, flags=re.DOTALL)
        text = re.sub(r'<ref[^>]*/>', '', text)
        text = re.sub(r'<[^>]+>', ' ', text)
        text = re.sub(r'^[ \t]*(?:\{\||\|\}|[|!]).+$', '', text, flags=re.MULTILINE)
        text = re.sub(r'^[*#:;]+\s*', '', text, flags=re.MULTILINE)
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()

    def _kowiki_to_docx(self, source: str, bz2_path: Path, docx_dir: Path) -> None:
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")

        ARTICLES_PER_FILE = 200
        MIN_TEXT_LEN = 300

        if list(docx_dir.glob("kowiki_*.docx")):
            print(f"kowiki: docx 파일이 이미 존재합니다. 건너뜁니다. ({docx_dir})")
            return

        batch_idx = 0
        count = 0
        doc = Document()
        title = ""
        ns_ok = True

        pbar = tqdm(desc="kowiki→docx", unit="article")
        try:
            with bz2.open(str(bz2_path), "rb") as f:
                for event, elem in ET.iterparse(f, events=("end",)):
                    tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag

                    if tag == "title":
                        title = elem.text or ""
                        elem.clear()
                    elif tag == "ns":
                        ns_ok = (elem.text or "0") == "0"
                        elem.clear()
                    elif tag == "text":
                        raw = elem.text or ""
                        elem.clear()
                        if not ns_ok or not title:
                            continue
                        if raw.startswith(("#REDIRECT", "#넘겨주기")):
                            continue
                        clean = self._strip_wiki_markup(raw)
                        if len(clean) < MIN_TEXT_LEN:
                            continue

                        doc.add_heading(title, level=1)
                        for para in clean.split("\n\n"):
                            para = para.strip()
                            if para:
                                doc.add_paragraph(para)
                        count += 1
                        pbar.update(1)

                        if count % ARTICLES_PER_FILE == 0:
                            fname = docx_dir / f"kowiki_{batch_idx:05d}.docx"
                            doc.save(str(fname))
                            self.log(source, str(bz2_path), fname, fname.stat().st_size, "docx_batch")
                            batch_idx += 1
                            doc = Document()
                            if not self.any_quota_remaining():
                                break
        finally:
            pbar.close()

        if count % ARTICLES_PER_FILE != 0:
            fname = docx_dir / f"kowiki_{batch_idx:05d}.docx"
            doc.save(str(fname))
            self.log(source, str(bz2_path), fname, fname.stat().st_size, "docx_batch")
            batch_idx += 1

        print(f"kowiki: {count}개 문서 → {batch_idx}개 docx 파일 ({docx_dir})")

    def run_common_crawl_ko(self) -> None:
        source = "common_crawl_ko_text"
        if not self.under_type_quota("dummy.txt"):
            return
        if ArchiveIterator is None:
            raise RuntimeError("warcio not installed. pip install -r requirements.txt")
        sconf = self.cfg.raw.get(source, {})
        crawl_id = sconf.get("crawl_id", "CC-MAIN-2026-12")
        max_wet = int(sconf.get("max_wet_files", 2000))
        min_ratio = float(sconf.get("min_hangul_ratio", 0.15))
        chunk_mb = int(sconf.get("output_chunk_mb", 128))
        base = f"https://data.commoncrawl.org/crawl-data/{crawl_id}/wet.paths.gz"
        paths_resp = self.get(base, stream=True)
        if not paths_resp:
            raise RuntimeError(f"Cannot fetch WET paths: {base}")
        paths = gzip.decompress(paths_resp.content).decode("utf-8", errors="ignore").splitlines()
        random.shuffle(paths)
        paths = paths[:max_wet]
        out_dir = self.cfg.root / source / "filtered_ko_txt"; mkdir(out_dir)
        raw_dir = self.cfg.root / source / "wet_raw"; mkdir(raw_dir)  # 임시 디렉토리 (WET 처리 후 즉시 삭제)
        chunk_target = chunk_mb * 1024 * 1024
        chunk_idx = len(list(out_dir.glob("ko_commoncrawl_*.txt")))
        out_path = out_dir / f"ko_commoncrawl_{chunk_idx:05d}.txt"
        out = out_path.open("ab")
        written_in_chunk = out_path.stat().st_size if out_path.exists() else 0
        pbar = tqdm(paths, desc=source, unit="wet")
        try:
            for rel in pbar:
                if not self.any_quota_remaining():
                    break
                if not self.under_type_quota("dummy.txt"):
                    break
                wet_url = "https://data.commoncrawl.org/" + rel
                r = self.get(wet_url, stream=True)
                if not r:
                    continue
                local = raw_dir / safe_name_from_url(wet_url, ".warc.wet.gz")
                with local.open("wb") as f:
                    shutil.copyfileobj(r.raw, f)
                quota_hit = False
                with gzip.open(local, "rb") as stream:
                    for record in ArchiveIterator(stream):
                        if not self.any_quota_remaining() or not self.under_type_quota("dummy.txt"):
                            quota_hit = True
                            break
                        if record.rec_type != "conversion":
                            continue
                        payload = record.content_stream().read().decode("utf-8", errors="ignore")
                        if len(payload) < 300:
                            continue
                        hangul = len(HANGUL_RE.findall(payload))
                        ratio = hangul / max(1, len(payload))
                        if ratio >= min_ratio:
                            block = ("\n\n---DOC---\n" + payload.strip() + "\n").encode("utf-8", errors="ignore")
                            out.write(block)
                            written_in_chunk += len(block)
                            if written_in_chunk >= chunk_target:
                                out.close()
                                self.log(source, "commoncrawl-ko-filtered", out_path, written_in_chunk, "ko_chunk")
                                chunk_idx += 1
                                out_path = out_dir / f"ko_commoncrawl_{chunk_idx:05d}.txt"
                                out = out_path.open("ab")
                                written_in_chunk = 0
                local.unlink(missing_ok=True)  # WET 원본 즉시 삭제
                if quota_hit:
                    break
                time.sleep(self.cfg.sleep)
        finally:
            out.close()

    def run_namuwiki_dump(self) -> None:
        source = "namuwiki_dump"
        if not self.under_type_quota("dummy.txt"):
            return
        sconf = self.cfg.raw.get(source, {})
        index_url = sconf.get("index_url", "https://mu-star.net/wikidb")
        chunk_mb = int(sconf.get("output_chunk_mb", 128))
        min_text_len = int(sconf.get("min_text_len", 300))

        # 덤프 인덱스 페이지에서 최신 다운로드 URL 찾기
        r = self.get(index_url)
        if not r:
            print(f"[{source}] 덤프 인덱스 페이지 접근 실패: {index_url}")
            return
        soup = BeautifulSoup(r.text, "lxml")
        dump_url = None
        for a in soup.find_all("a", href=True):
            href = a["href"]
            if re.search(r'\.(json\.gz|json\.zst|db\.gz|db\.zst|json)$', href, re.I):
                dump_url = urljoin(index_url, href)
                break
        if not dump_url:
            print(f"[{source}] 덤프 파일 링크를 찾지 못했습니다. 수동으로 dump_url을 config에 지정하세요.")
            return

        out_dir = self.cfg.root / source / "txt"
        mkdir(out_dir)
        raw_dir = self.cfg.root / source / "raw"
        mkdir(raw_dir)

        local = raw_dir / safe_name_from_url(dump_url)
        if not local.exists():
            print(f"[{source}] 덤프 다운로드 중: {dump_url}")
            r2 = self.get(dump_url, stream=True)
            if not r2:
                print(f"[{source}] 덤프 다운로드 실패")
                return
            with local.open("wb") as f:
                shutil.copyfileobj(r2.raw, f)
            print(f"[{source}] 덤프 다운로드 완료: {local}")

        chunk_target = chunk_mb * 1024 * 1024
        chunk_idx = len(list(out_dir.glob("namuwiki_*.txt")))
        out_path = out_dir / f"namuwiki_{chunk_idx:05d}.txt"
        out = out_path.open("ab")
        written_in_chunk = out_path.stat().st_size if out_path.exists() else 0

        import json as _json
        opener = gzip.open if str(local).endswith(".gz") else open
        pbar = tqdm(desc=source, unit="article")
        try:
            with opener(local, "rt", encoding="utf-8", errors="ignore") as f:
                data = _json.load(f)
            for entry in data:
                if not self.under_type_quota("dummy.txt"):
                    break
                ns = entry.get("namespace", 0)
                if ns != 0:
                    continue
                text = entry.get("text", "") or ""
                text = self._strip_namu_markup(text)
                if len(text) < min_text_len:
                    continue
                title = entry.get("title", "")
                block = f"\n\n---DOC---\n제목: {title}\n{text.strip()}\n".encode("utf-8", errors="ignore")
                out.write(block)
                written_in_chunk += len(block)
                pbar.update(1)
                if written_in_chunk >= chunk_target:
                    out.close()
                    self.log(source, dump_url, out_path, written_in_chunk, "ko_chunk")
                    chunk_idx += 1
                    out_path = out_dir / f"namuwiki_{chunk_idx:05d}.txt"
                    out = out_path.open("ab")
                    written_in_chunk = 0
        except Exception as e:
            print(f"[{source}] 파싱 오류: {e}")
        finally:
            out.close()
            pbar.close()
            if out_path.exists() and out_path.stat().st_size > 0:
                self.log(source, dump_url, out_path, written_in_chunk, "ko_chunk")

    @staticmethod
    def _strip_namu_markup(text: str) -> str:
        text = re.sub(r'\[{2}[^|\]]*\|([^\]]*)\]{2}', r'\1', text)  # [[링크|표시]] → 표시
        text = re.sub(r'\[{2}([^\]]*)\]{2}', r'\1', text)            # [[링크]] → 링크
        text = re.sub(r'\{{3}[^|]*\|([^}]*)\}{3}', r'\1', text)      # {{{색|텍스트}}} → 텍스트
        text = re.sub(r'\{{3}.*?\}{3}', '', text, flags=re.DOTALL)    # {{{...}}} 제거
        text = re.sub(r'##.*$', '', text, flags=re.MULTILINE)         # 주석 제거
        text = re.sub(r'\[(?:include|youtube|anchor|목차|tableofcontents)[^\]]*\]', '', text, flags=re.I)
        text = re.sub(r'\[(?:각주|footnote)[^\]]*\]', '', text, flags=re.I)
        text = re.sub(r'\|\|[^\n]*', '', text)                        # 테이블 제거
        text = re.sub(r"'{2,}", '', text)                             # 볼드/이탤릭 마커
        text = re.sub(r'={2,6}\s*(.+?)\s*={2,6}', r'\n\1\n', text)  # 제목
        text = re.sub(r'<[^>]+>', ' ', text)                          # HTML 태그
        text = re.sub(r'https?://\S+', '', text)                      # URL
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()

    def run_law_lbook(self) -> None:
        """국가법령정보센터 전자법령집(lbook) — 목록 → 상세 → 실제 PDF 다운로드.
        HWP는 HWPML(구형 XML) 포맷이라 LibreOffice 변환이 되지 않아 제외 — 같은 책의 PDF로 내용 커버됨."""
        source = "law_lbook"
        sconf = self.cfg.raw.get(source, {})
        base = "https://www.law.go.kr/lbook/"
        list_url_tpl = sconf.get(
            "list_url_tpl",
            "https://www.law.go.kr/lbook/lbListR.do?dSearchYn=N&FSort=100"
            "&pageIndex={page}&menuId=13&subMenuId=67&tabMenuId=293",
        )
        max_pages = int(sconf.get("max_pages", 90))
        out_dir = self.cfg.root / source / "files"
        mkdir(out_dir)

        seq_re = re.compile(r"lbInfoR\.do\?lbookSeq=(\d+)")
        file_re = re.compile(r'href="(lbFileDownload\.do\?[^"]+)"', re.I)

        seqs: List[str] = []
        seen_seq: Set[str] = set()
        pbar = tqdm(range(1, max_pages + 1), desc=f"{source}-list", unit="page")
        for page in pbar:
            if not self.any_quota_remaining():
                break
            r = self.get(list_url_tpl.format(page=page))
            if not r:
                continue
            for m in seq_re.finditer(r.text):
                seq = m.group(1)
                if seq not in seen_seq:
                    seen_seq.add(seq)
                    seqs.append(seq)
            time.sleep(self.cfg.sleep)
        pbar.close()

        pbar2 = tqdm(seqs, desc=f"{source}-files", unit="book")
        for seq in pbar2:
            if not self.any_quota_remaining():
                break
            r = self.get(urljoin(base, f"lbInfoR.do?lbookSeq={seq}"))
            if not r:
                continue
            for m in file_re.finditer(r.text):
                href = m.group(1).replace("&amp;", "&")
                if "flext=pdf" not in href.lower():
                    continue  # hwp(HWPML, 변환 불가)/zip 등은 제외 — pdf만 수집
                self.download(source, urljoin(base, href), out_dir)
            time.sleep(self.cfg.sleep)
        pbar2.close()

    def run_source(self, source: str) -> None:
        if source == "kowiki_knowledge":
            self.run_kowiki(); return
        if source == "common_crawl_ko_text":
            self.run_common_crawl_ko(); return
        if source == "law_lbook":
            self.run_law_lbook(); return
        if source == "namuwiki_dump":
            self.run_namuwiki_dump(); return
        conf = self.cfg.raw.get(source, {})
        seeds = conf.get("seed_urls", [])
        if not seeds:
            print(f"No seed URLs for {source}")
            return
        self.crawl_generic(source, seeds)

    def plan(self) -> None:
        tq = self.cfg.type_quotas
        target_sum = sum(tq.values())
        # 타입별로 목표치를 초과한 분량은 진행률 계산에서 제외(캡)한다 — 한 타입의
        # 초과 수집이 전체 진행률을 부풀려 보이게 하는 것을 방지 (2026-07-03: txt 초과 건).
        capped_sum = sum(min(self._type_bytes.get(g, 0), limit) for g, limit in tq.items())
        print(f"\nRoot : {self.cfg.root}")
        print(f"수집량(원본, 초과분 포함): {human(self.total_collected())}")
        print(f"목표 합계(타입별 쿼터 합): {human(target_sum)}")
        print()
        lang_caps = self.cfg.lang_caps
        print(f"{'Type group':<18} {'Collected':>10} {'Target':>10} {'Remaining':>10}  Progress")
        print("-" * 68)
        for grp in TYPE_GROUPS:
            limit = tq.get(grp, 0)
            cur = self._type_bytes.get(grp, 0)
            rem = max(0, limit - cur)
            pct = min(100.0, cur / limit * 100) if limit else 0.0
            bar = "#" * int(pct / 5) + "." * (20 - int(pct / 5))
            line = f"{grp:<18} {human(cur):>10} {human(limit):>10} {human(rem):>10}  [{bar}] {pct:5.1f}%"
            if cur > limit:
                line += f"   (초과 +{human(cur - limit)})"
            if grp in lang_caps:
                en_cur = self._type_en_bytes.get(grp, 0)
                en_pct = (en_cur / cur * 100) if cur else 0.0
                line += f"   (en {human(en_cur)}, {en_pct:.0f}% / cap {lang_caps[grp]*100:.0f}%)"
            print(line)
        print()
        overall_pct = min(100.0, capped_sum / target_sum * 100) if target_sum else 0.0
        print(f"Overall progress: {overall_pct:.1f}%  (타입별 목표 기준, 초과분 제외)")


def _exec_sources(col: "Collector", sources: List[str], workers: int) -> None:
    """소스 목록을 순차 또는 병렬로 실행."""
    if workers > 1:
        from concurrent.futures import ThreadPoolExecutor

        def _run_one(src: str) -> None:
            if not col.any_quota_remaining():
                return
            print(f"\n=== Running {src} ===", flush=True)
            col.run_source(src)

        print(f"[병렬 수집] workers={workers}, sources={len(sources)}", flush=True)
        with ThreadPoolExecutor(max_workers=workers) as pool:
            pool.map(_run_one, sources)
    else:
        for src in sources:
            if not col.any_quota_remaining():
                print("All type quotas filled. Done.")
                break
            print(f"\n=== Running {src} ===")
            col.run_source(src)
            col.plan()


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--config", default="config.yaml")
    ap.add_argument("--plan", action="store_true")
    ap.add_argument("--run", nargs="+", help="source names or 'all'")
    ap.add_argument("--parallel", type=int, default=0,
                    help="parallel source workers (0 = use config parallel_workers)")
    ap.add_argument("--bg", action="store_true", help="run in background (detached from terminal)")
    ap.add_argument("--log", default="collect.log", help="log file path when using --bg")
    args = ap.parse_args()

    if args.bg:
        cmd = [sys.executable] + [a for a in sys.argv[1:] if a not in ("--bg",)]
        log_path = Path(args.log)
        with log_path.open("a") as f:
            proc = subprocess.Popen(cmd, stdout=f, stderr=f, start_new_session=True)
        print(f"Background PID: {proc.pid}  log: {log_path.resolve()}")
        print(f"  tail -f {log_path}")
        print(f"  kill {proc.pid}")
        return

    cfg = Cfg.load(args.config)
    col = Collector(cfg)
    if args.plan:
        col.plan()
    if args.run:
        all_cfg = cfg.all_sources() if "all" in args.run else args.run
        workers = args.parallel if args.parallel > 0 else cfg.parallel_workers

        # 일반 소스: enabled이고 auto_pool이 아닌 것
        regular = [s for s in all_cfg
                   if cfg.raw.get(s, {}).get("enabled", True)
                   and not cfg.raw.get(s, {}).get("auto_pool")]
        # 풀 소스: auto_pool: true — 타입 미달 시 자동 추가
        pool_srcs = [s for s in all_cfg if cfg.raw.get(s, {}).get("auto_pool")]

        used_pool: set = set()
        to_run = regular

        while to_run:
            _exec_sources(col, to_run, workers)
            col.plan()

            if not col.any_quota_remaining():
                break

            lagging = col.lagging_types()
            promoted = [s for s in pool_srcs
                        if s not in used_pool
                        and cfg.raw.get(s, {}).get("type_hint") in lagging]
            if not promoted:
                break

            print(f"\n[자동 소스 추가] 미달 타입={sorted(lagging)}", flush=True)
            for s in promoted:
                print(f"  + {s}  (type_hint={cfg.raw[s].get('type_hint')})", flush=True)
            used_pool.update(promoted)
            to_run = promoted

if __name__ == "__main__":
    main()
